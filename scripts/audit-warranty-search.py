# -*- coding: utf-8 -*-
import json, re
from pathlib import Path
from docx import Document
from pypdf import PdfReader

DIR = Path(r"c:\Users\Windows\Desktop\DOCUMENTACIONES GENERALES EMPRENOR")
OUT = Path(r"c:\Users\Windows\Downloads\emprenor_app\nextjs_space\docs\emprenor-warranty-search.json")

def read_docx(p):
    doc = Document(p)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for row in t.rows:
            parts.extend([c.text for c in row.cells if c.text.strip()])
    return "\n".join(parts)

def read_pdf(p, max_pages=80):
    r = PdfReader(str(p))
    return "\n".join((page.extract_text() or "") for page in r.pages[:max_pages])

patterns = [
    (r"120\s*d[ií]as", "120_dias"),
    (r"garant[ií]a.{0,80}", "garantia"),
    (r"post\s*[- ]?venta.{0,60}", "postventa"),
    (r"servicio\s+de\s+calidad.{0,60}", "servicio_calidad"),
    (r"CONF-EL-001", "conf_el"),
    (r"recepci[oó]n\s+final", "recepcion_final"),
    (r"control\s+document", "control_doc"),
    (r"matriz\s+maestra", "matriz_maestra"),
    (r"retenci[oó]n.{0,40}", "retencion"),
]

results = {}
for f in sorted(DIR.iterdir()):
    if f.suffix.lower() not in (".docx", ".pdf"):
        continue
    try:
        text = read_docx(f) if f.suffix.lower() == ".docx" else read_pdf(f)
    except Exception as e:
        results[f.name] = {"error": str(e)}
        continue
    hits = {}
    for pat, key in patterns:
        found = []
        for m in re.finditer(pat, text, re.I | re.DOTALL):
            snip = re.sub(r"\s+", " ", text[max(0, m.start() - 80) : m.end() + 80]).strip()
            if snip not in found:
                found.append(snip[:300])
        if found:
            hits[key] = found[:5]
    results[f.name] = hits

OUT.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
print("done", sum(1 for v in results.values() if v))
