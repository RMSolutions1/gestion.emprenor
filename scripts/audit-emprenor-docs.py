# -*- coding: utf-8 -*-
import json, re, sys
from pathlib import Path
from docx import Document
from pypdf import PdfReader

DIR = Path(r"c:\Users\Windows\Desktop\DOCUMENTACIONES GENERALES EMPRENOR")
OUT = Path(r"c:\Users\Windows\Downloads\emprenor_app\nextjs_space\docs\emprenor-docs-audit.json")

def read_docx(p):
    doc = Document(p)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for row in t.rows:
            parts.extend([c.text.strip() for c in row.cells if c.text.strip()])
    return "\n".join(parts)

def read_pdf(p, max_pages=60):
    r = PdfReader(str(p))
    return "\n".join((page.extract_text() or "") for page in r.pages[:max_pages])

def extract_meta(text, fname):
    # code patterns
    codes = list(set(re.findall(r"\b(PAC-[A-Z]{2}-\d{3}|PS-EL-\d{3}|SIGCE[^\s,]*)", text, re.I)))
    rev = re.findall(r"Rev\.?\s*(\d+)|Revisi[oó]n\s*:?\s*(\d+)", text, re.I)
    revs = [a or b for a, b in rev[:3]]

    # warranty / 120 days
    warranty = []
    for m in re.finditer(
        r".{0,120}(120\s*d[ií]as|garant[ií]a\s+de\s+\d+|garant[ií]a\s+m[ií]nima|post\s*[- ]?venta|servicio\s+post).{0,120}",
        text,
        re.I | re.DOTALL,
    ):
        snippet = re.sub(r"\s+", " ", m.group(0)).strip()
        if len(snippet) > 30:
            warranty.append(snippet[:280])

    # document control sections
    doc_ctrl = bool(re.search(r"control\s+document|gesti[oó]n\s+document|trazabilidad|retenci[oó]n|archivo", text, re.I))

    # ISO mapping
    iso = bool(re.search(r"ISO\s*9001", text, re.I))
    sigce = bool(re.search(r"SIGCE", text, re.I))
    aea = len(re.findall(r"90364|AEA", text, re.I))

    # index / sections
    sections = re.findall(r"^\s*(\d+\.[\d\.]*\s+[A-ZÁÉÍÓÚÑ][^\n]{5,80})", text, re.M)[:20]

    return {
        "codes": codes,
        "revisions": revs,
        "warranty_snippets": warranty[:8],
        "has_doc_management": doc_ctrl,
        "mentions_iso9001": iso,
        "mentions_sigce": sigce,
        "aea_refs_count": aea,
        "section_headings": sections[:12],
        "char_count": len(text),
    }

results = []
for f in sorted(DIR.iterdir()):
    if f.suffix.lower() not in (".docx", ".pdf"):
        continue
    entry = {"file": f.name, "format": f.suffix.lower(), "size_kb": round(f.stat().st_size / 1024, 1)}
    try:
        text = read_docx(f) if f.suffix.lower() == ".docx" else read_pdf(f)
        entry["readable"] = len(text.strip()) > 200
        entry["first_lines"] = [l.strip() for l in text.split("\n") if l.strip()][:10]
        entry.update(extract_meta(text, f.name))
    except Exception as e:
        entry["error"] = str(e)
        entry["readable"] = False
    results.append(entry)

OUT.parent.mkdir(parents=True, exist_ok=True)
OUT.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
print(f"Wrote {OUT} ({len(results)} files)")
